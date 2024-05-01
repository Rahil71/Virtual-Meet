from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def add_title(document, title):
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = title
    header_paragraph.style.font.size = Pt(20)
    header_paragraph.style.font.name = 'Calibri'
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_summary(document, summary):
    summary_paragraph = document.add_paragraph(summary)
    summary_paragraph.style.font.size = Pt(13)
    summary_paragraph.style.font.name = 'Times New Roman'
    summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def add_diarization(document, diarization):
    diarization_paragraph = document.add_paragraph(diarization)
    diarization_paragraph.style.font.size = Pt(13)
    diarization_paragraph.style.font.name = 'Times New Roman'




def create_document(title,summary,diarization):
    document = Document()

    # Adding title to header
    add_title(document, title)

    # Adding summary section
    document.add_heading('Summary', level=1)
    add_summary(document, summary)

    # Adding diarization section
    document.add_heading('Segmentation', level=1)
    add_diarization(document, diarization)
    print('came_here')
    # Save the document
    document.save('meeting_details.docx')
